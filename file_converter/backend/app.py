"""
File Converter & Merger — Flask Backend
Run: python app.py
API base: http://localhost:5000/api
"""

import io
import os
import csv
import zipfile
import tempfile
from pathlib import Path
from flask import Flask, request, jsonify, send_file
from flask import make_response

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50 MB

# ── CORS (manual, no flask-cors needed) ──────────────────────────────────────
@app.after_request
def add_cors(resp):
    resp.headers["Access-Control-Allow-Origin"] = "*"
    resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
    resp.headers["Access-Control-Allow-Methods"] = "GET,POST,OPTIONS"
    return resp

@app.route("/api/<path:p>", methods=["OPTIONS"])
def options_handler(p):
    return make_response("", 204)


# ── Conversion matrix ─────────────────────────────────────────────────────────
# Maps source_ext -> list of valid target_exts
CONVERSION_MAP = {
    "pdf":  ["txt", "png", "jpg", "jpeg"],
    "csv":  ["xlsx", "txt", "pdf"],
    "xlsx": ["csv", "txt", "pdf"],
    "doc":  ["txt", "pdf"],
    "txt":  ["pdf", "csv", "xlsx", "doc"],
    "png":  ["jpg", "jpeg", "pdf"],
    "jpg":  ["png", "jpeg", "pdf"],
    "jpeg": ["png", "jpg", "pdf"],
}

MERGEABLE = {
    "pdf", "csv", "xlsx", "txt", "png", "jpg", "jpeg", "doc"
}

ALLOWED_EXT = set(CONVERSION_MAP.keys())


def ext(filename: str) -> str:
    return Path(filename).suffix.lstrip(".").lower()


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route("/api/conversion-options", methods=["POST"])
def conversion_options():
    """Return allowed target formats for the uploaded file(s)."""
    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "No files uploaded"}), 400

    exts = list({ext(f.filename) for f in files})
    invalid = [e for e in exts if e not in ALLOWED_EXT]
    if invalid:
        return jsonify({"error": f"Unsupported file type(s): {', '.join(invalid)}"}), 400

    if len(exts) > 1:
        # multiple different extensions → only common targets
        targets = set(CONVERSION_MAP[exts[0]])
        for e in exts[1:]:
            targets &= set(CONVERSION_MAP[e])
        targets -= set(exts)           # can't convert to same type
        return jsonify({"source_exts": exts, "targets": sorted(targets)})

    source = exts[0]
    targets = [t for t in CONVERSION_MAP[source] if t != source]
    return jsonify({"source_exts": [source], "targets": targets})


@app.route("/api/convert", methods=["POST"])
def convert():
    files  = request.files.getlist("files")
    target = request.form.get("target", "").lower().lstrip(".")

    if not files or not target:
        return jsonify({"error": "Missing files or target format"}), 400

    results = []
    errors  = []

    for f in files:
        src_ext = ext(f.filename)
        if src_ext == target:
            errors.append(f"{f.filename}: source and target format are the same ({target})")
            continue
        if target not in CONVERSION_MAP.get(src_ext, []):
            errors.append(f"{f.filename}: cannot convert .{src_ext} → .{target}")
            continue
        try:
            data, mime = _do_convert(f, src_ext, target)
            stem = Path(f.filename).stem
            results.append((f"{stem}.{target}", mime, data))
        except Exception as e:
            errors.append(f"{f.filename}: {e}")

    if errors and not results:
        return jsonify({"error": "; ".join(errors)}), 422

    if len(results) == 1:
        name, mime, data = results[0]
        buf = io.BytesIO(data)
        buf.seek(0)
        resp = send_file(buf, mimetype=mime, as_attachment=True, download_name=name)
        if errors:
            resp.headers["X-Warnings"] = " | ".join(errors)
        return resp

    # multiple → zip
    zbuf = io.BytesIO()
    with zipfile.ZipFile(zbuf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, _, data in results:
            zf.writestr(name, data)
    zbuf.seek(0)
    return send_file(zbuf, mimetype="application/zip",
                     as_attachment=True, download_name="converted_files.zip")


@app.route("/api/merge", methods=["POST"])
def merge():
    files = request.files.getlist("files")
    if not files or len(files) < 2:
        return jsonify({"error": "Upload at least 2 files to merge"}), 400

    exts = {ext(f.filename) for f in files}
    if len(exts) > 1:
        return jsonify({
            "error": f"All files must have the same extension to merge. "
                     f"Found: {', '.join(sorted(exts))}"
        }), 422

    src_ext = exts.pop()
    if src_ext not in MERGEABLE:
        return jsonify({"error": f"Merging .{src_ext} files is not supported"}), 422

    try:
        data, mime, out_name = _do_merge(files, src_ext)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    buf = io.BytesIO(data)
    buf.seek(0)
    return send_file(buf, mimetype=mime, as_attachment=True, download_name=out_name)


# ── Conversion logic ──────────────────────────────────────────────────────────

def _do_convert(file_storage, src_ext: str, target: str) -> tuple[bytes, str]:
    raw = file_storage.read()

    # ── image conversions ────────────────────────────────────────────────────
    if src_ext in ("png", "jpg", "jpeg") and target in ("png", "jpg", "jpeg"):
        from PIL import Image
        img = Image.open(io.BytesIO(raw)).convert("RGB")
        fmt = "JPEG" if target in ("jpg", "jpeg") else "PNG"
        out = io.BytesIO()
        img.save(out, format=fmt)
        mime = "image/jpeg" if fmt == "JPEG" else "image/png"
        return out.getvalue(), mime

    if src_ext in ("png", "jpg", "jpeg") and target == "pdf":
        from PIL import Image
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Image as RLImage
        img = Image.open(io.BytesIO(raw))
        w, h = img.size
        tmp_img = tempfile.NamedTemporaryFile(suffix=f".{src_ext}", delete=False)
        tmp_img.write(raw); tmp_img.flush()
        out = io.BytesIO()
        doc = SimpleDocTemplate(out, pagesize=letter)
        pw, ph = letter
        ratio = min(pw / w, ph / h) * 0.9
        doc.build([RLImage(tmp_img.name, width=w*ratio, height=h*ratio)])
        os.unlink(tmp_img.name)
        return out.getvalue(), "application/pdf"

    if src_ext == "pdf" and target == "txt":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(raw))
        text = "\n\n".join(p.extract_text() or "" for p in reader.pages)
        return text.encode("utf-8"), "text/plain"

    if src_ext == "pdf" and target in ("png", "jpg", "jpeg"):
        # Convert first page via pypdf text → placeholder image
        from pypdf import PdfReader
        from reportlab.pdfgen import canvas as rl_canvas
        from PIL import Image as PILImage
        reader = PdfReader(io.BytesIO(raw))
        text = (reader.pages[0].extract_text() or "PDF page (no text layer)").strip()[:500]
        # Draw text onto a white canvas, save as image
        tmp = io.BytesIO()
        c = rl_canvas.Canvas(tmp, pagesize=(600, 800))
        c.setFont("Helvetica", 11)
        y = 770
        for line in text.split("\n"):
            c.drawString(30, y, line[:90])
            y -= 16
            if y < 30:
                break
        c.save()
        tmp.seek(0)
        img = PILImage.open(tmp)
        out = io.BytesIO()
        fmt = "JPEG" if target in ("jpg", "jpeg") else "PNG"
        img.save(out, format=fmt)
        return out.getvalue(), "image/jpeg" if fmt == "JPEG" else "image/png"

    if src_ext == "csv" and target == "xlsx":
        import pandas as pd
        df = pd.read_csv(io.BytesIO(raw))
        out = io.BytesIO()
        df.to_excel(out, index=False)
        return out.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    if src_ext == "csv" and target == "txt":
        return raw, "text/plain"

    if src_ext == "csv" and target == "pdf":
        import pandas as pd
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet
        df = pd.read_csv(io.BytesIO(raw))
        out = io.BytesIO()
        doc = SimpleDocTemplate(out, pagesize=letter)
        styles = getSampleStyleSheet()
        data = [list(df.columns)] + df.fillna("").astype(str).values.tolist()
        t = Table(data)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",   (0, 0), (-1, -1), 8),
            ("GRID",       (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EBF0FA")]),
        ]))
        doc.build([Paragraph("Converted from CSV", styles["Heading2"]), t])
        return out.getvalue(), "application/pdf"

    if src_ext == "xlsx" and target == "csv":
        import pandas as pd
        df = pd.read_excel(io.BytesIO(raw))
        out = io.BytesIO()
        df.to_csv(out, index=False)
        return out.getvalue(), "text/csv"

    if src_ext == "xlsx" and target == "txt":
        import pandas as pd
        df = pd.read_excel(io.BytesIO(raw))
        return df.to_string(index=False).encode("utf-8"), "text/plain"

    if src_ext == "xlsx" and target == "pdf":
        import pandas as pd
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet
        df = pd.read_excel(io.BytesIO(raw))
        out = io.BytesIO()
        doc = SimpleDocTemplate(out, pagesize=letter)
        styles = getSampleStyleSheet()
        data = [list(df.columns)] + df.fillna("").astype(str).values.tolist()
        t = Table(data)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#70AD47")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",   (0, 0), (-1, -1), 8),
            ("GRID",       (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F0F7EA")]),
        ]))
        doc.build([Paragraph("Converted from XLSX", styles["Heading2"]), t])
        return out.getvalue(), "application/pdf"

    if src_ext == "txt" and target == "pdf":
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        text = raw.decode("utf-8", errors="replace")
        out = io.BytesIO()
        doc = SimpleDocTemplate(out, pagesize=letter)
        styles = getSampleStyleSheet()
        paras = [Paragraph(line.replace("&", "&amp;").replace("<", "&lt;") or "&nbsp;",
                           styles["Normal"])
                 for line in text.split("\n")]
        doc.build(paras)
        return out.getvalue(), "application/pdf"

    if src_ext == "txt" and target == "csv":
        lines = raw.decode("utf-8", errors="replace").splitlines()
        out = io.BytesIO()
        w = csv.writer(io.StringIO())
        rows = [{"line_no": i+1, "content": l} for i, l in enumerate(lines)]
        out_str = io.StringIO()
        writer = csv.DictWriter(out_str, fieldnames=["line_no", "content"])
        writer.writeheader(); writer.writerows(rows)
        return out_str.getvalue().encode("utf-8"), "text/csv"

    if src_ext == "txt" and target == "xlsx":
        import pandas as pd
        lines = raw.decode("utf-8", errors="replace").splitlines()
        df = pd.DataFrame({"line_no": range(1, len(lines)+1), "content": lines})
        out = io.BytesIO()
        df.to_excel(out, index=False)
        return out.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    if src_ext == "txt" and target == "doc":
        from docx import Document
        doc = Document()
        text = raw.decode("utf-8", errors="replace")
        for line in text.split("\n"):
            doc.add_paragraph(line)
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    if src_ext == "doc" and target == "txt":
        from docx import Document
        doc = Document(io.BytesIO(raw))
        text = "\n".join(p.text for p in doc.paragraphs)
        return text.encode("utf-8"), "text/plain"

    if src_ext == "doc" and target == "pdf":
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        doc = Document(io.BytesIO(raw))
        text = "\n".join(p.text for p in doc.paragraphs)
        out = io.BytesIO()
        rdoc = SimpleDocTemplate(out, pagesize=letter)
        styles = getSampleStyleSheet()
        paras = [Paragraph(line.replace("&", "&amp;").replace("<", "&lt;") or "&nbsp;",
                           styles["Normal"])
                 for line in text.split("\n")]
        rdoc.build(paras)
        return out.getvalue(), "application/pdf"

    raise ValueError(f"Conversion .{src_ext} → .{target} not implemented")


# ── Merge logic ───────────────────────────────────────────────────────────────

def _do_merge(files, src_ext: str) -> tuple[bytes, str, str]:

    if src_ext == "pdf":
        from pypdf import PdfWriter
        writer = PdfWriter()
        for f in files:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(f.read()))
            for page in reader.pages:
                writer.add_page(page)
        out = io.BytesIO()
        writer.write(out)
        return out.getvalue(), "application/pdf", "merged.pdf"

    if src_ext in ("csv",):
        import pandas as pd
        dfs = [pd.read_csv(io.BytesIO(f.read())) for f in files]
        merged = pd.concat(dfs, ignore_index=True)
        out = io.BytesIO()
        merged.to_csv(out, index=False)
        return out.getvalue(), "text/csv", "merged.csv"

    if src_ext == "xlsx":
        import pandas as pd
        dfs = [pd.read_excel(io.BytesIO(f.read())) for f in files]
        merged = pd.concat(dfs, ignore_index=True)
        out = io.BytesIO()
        merged.to_excel(out, index=False)
        return out.getvalue(), \
               "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", \
               "merged.xlsx"

    if src_ext == "txt":
        parts = []
        for f in files:
            parts.append(f.read().decode("utf-8", errors="replace"))
        merged = "\n\n--- merged ---\n\n".join(parts)
        return merged.encode("utf-8"), "text/plain", "merged.txt"

    if src_ext == "doc":
        from docx import Document
        merged_doc = Document()
        for i, f in enumerate(files):
            doc = Document(io.BytesIO(f.read()))
            if i > 0:
                merged_doc.add_page_break()
            for para in doc.paragraphs:
                merged_doc.add_paragraph(para.text, para.style.name if para.style else None)
        out = io.BytesIO()
        merged_doc.save(out)
        return out.getvalue(), \
               "application/vnd.openxmlformats-officedocument.wordprocessingml.document", \
               "merged.docx"

    if src_ext in ("png", "jpg", "jpeg"):
        from PIL import Image
        images = [Image.open(io.BytesIO(f.read())).convert("RGB") for f in files]
        total_h = sum(img.height for img in images)
        max_w   = max(img.width  for img in images)
        canvas  = Image.new("RGB", (max_w, total_h), (255, 255, 255))
        y = 0
        for img in images:
            canvas.paste(img, (0, y)); y += img.height
        out = io.BytesIO()
        fmt  = "JPEG" if src_ext in ("jpg", "jpeg") else "PNG"
        mime = "image/jpeg" if fmt == "JPEG" else "image/png"
        canvas.save(out, format=fmt)
        return out.getvalue(), mime, f"merged.{src_ext}"

    raise ValueError(f"Merging .{src_ext} not supported")


if __name__ == "__main__":
    print("\n🔄  File Converter API")
    print("   Running on http://localhost:5000")
    print("   Press Ctrl+C to stop\n")
    app.run(debug=True, port=5000)
